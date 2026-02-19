# ============================================
# processors/word_processor.py
# Extract text from Word and Excel files
# ============================================

import docx
import pandas as pd


def extract_text_from_word(file_path: str) -> str:
    try:
        doc = docx.Document(file_path)
        full_text = []

        for para in doc.paragraphs:
            if para.text.strip():
                full_text.append(para.text.strip())

        for table in doc.tables:
            for row in table.rows:
                row_text = " | ".join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    full_text.append(row_text)

        return "\n".join(full_text) if full_text else "[No text found in Word document]"

    except Exception as e:
        return f"[ERROR reading Word file: {str(e)}]"


def extract_text_from_excel(file_path: str) -> str:
    try:
        xl = pd.ExcelFile(file_path)
        full_text = []

        for sheet_name in xl.sheet_names:
            df = xl.parse(sheet_name)
            full_text.append(f"\n--- Sheet: {sheet_name} ---")
            full_text.append(df.to_string(index=False))

        return "\n".join(full_text) if full_text else "[No data found in Excel file]"

    except Exception as e:
        return f"[ERROR reading Excel file: {str(e)}]"